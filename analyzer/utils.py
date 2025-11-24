"""
Utility functions for parsing different file types.

This module contains functions to extract text content from various file formats
including PDF, Word documents, Excel spreadsheets, and emails.
"""
import os
import re
import chardet
from langdetect import detect
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
import email
from email import policy
from email.parser import BytesParser


def detect_language(text):
    """
    Detect the language of the given text.
    
    Args:
        text: Text string to analyze
        
    Returns:
        str: Language code (e.g., 'en', 'ko')
    """
    try:
        if text and len(text.strip()) > 10:
            return detect(text)
        return 'unknown'
    except Exception:
        return 'unknown'


def parse_pdf(file_path):
    """
    Extract text content from a PDF file.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        str: Extracted text content
    """
    try:
        reader = PdfReader(file_path)
        text_content = []
        
        # Extract text from each page
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_content.append(text)
        
        return '\n\n'.join(text_content)
    except Exception as e:
        return f"Error parsing PDF: {str(e)}"


def parse_docx(file_path):
    """
    Extract text content from a Word document (.docx).
    
    Args:
        file_path: Path to the .docx file
        
    Returns:
        str: Extracted text content
    """
    try:
        doc = DocxDocument(file_path)
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(' | '.join(row_text))
        
        return '\n\n'.join(text_content)
    except Exception as e:
        return f"Error parsing DOCX: {str(e)}"


def parse_xlsx(file_path):
    """
    Extract text content from an Excel spreadsheet (.xlsx).
    
    Args:
        file_path: Path to the .xlsx file
        
    Returns:
        str: Extracted text content
    """
    try:
        workbook = load_workbook(file_path, data_only=True)
        text_content = []
        
        # Extract text from each sheet
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = [f"Sheet: {sheet_name}"]
            
            # Extract data from cells
            for row in sheet.iter_rows(values_only=True):
                row_data = []
                for cell_value in row:
                    if cell_value is not None:
                        row_data.append(str(cell_value))
                if row_data:
                    sheet_data.append(' | '.join(row_data))
            
            if len(sheet_data) > 1:
                text_content.append('\n'.join(sheet_data))
        
        return '\n\n'.join(text_content)
    except Exception as e:
        return f"Error parsing XLSX: {str(e)}"


def parse_email(file_path):
    """
    Extract text content from an email file (.eml or .msg).
    
    Args:
        file_path: Path to the email file
        
    Returns:
        str: Extracted text content
    """
    try:
        with open(file_path, 'rb') as f:
            msg = BytesParser(policy=policy.default).parse(f)
        
        email_content = []
        
        # Extract headers
        email_content.append(f"From: {msg.get('From', 'Unknown')}")
        email_content.append(f"To: {msg.get('To', 'Unknown')}")
        email_content.append(f"Subject: {msg.get('Subject', 'No Subject')}")
        email_content.append(f"Date: {msg.get('Date', 'Unknown')}")
        email_content.append("")
        
        # Extract body
        if msg.is_multipart():
            # Handle multipart messages
            for part in msg.walk():
                content_type = part.get_content_type()
                if content_type == 'text/plain':
                    payload = part.get_payload(decode=True)
                    if payload:
                        # Try to detect encoding
                        encoding = chardet.detect(payload)['encoding'] or 'utf-8'
                        try:
                            email_content.append(payload.decode(encoding))
                        except UnicodeDecodeError:
                            email_content.append(payload.decode('utf-8', errors='ignore'))
                elif content_type == 'text/html':
                    # For HTML emails, we'll extract text (basic)
                    payload = part.get_payload(decode=True)
                    if payload:
                        encoding = chardet.detect(payload)['encoding'] or 'utf-8'
                        try:
                            html_content = payload.decode(encoding)
                            # Simple HTML tag removal (basic implementation)
                            text_content = re.sub('<[^<]+?>', '', html_content)
                            email_content.append(text_content)
                        except UnicodeDecodeError:
                            pass
        else:
            # Handle simple messages
            payload = msg.get_payload(decode=True)
            if payload:
                encoding = chardet.detect(payload)['encoding'] or 'utf-8'
                try:
                    email_content.append(payload.decode(encoding))
                except UnicodeDecodeError:
                    email_content.append(payload.decode('utf-8', errors='ignore'))
        
        return '\n'.join(email_content)
    except Exception as e:
        return f"Error parsing email: {str(e)}"


def parse_file(file_path, file_type):
    """
    Parse a file based on its type.
    
    Args:
        file_path: Path to the file
        file_type: Type of file ('pdf', 'docx', 'xlsx', 'email')
        
    Returns:
        str: Extracted text content
    """
    if file_type == 'pdf':
        return parse_pdf(file_path)
    elif file_type == 'docx':
        return parse_docx(file_path)
    elif file_type == 'xlsx':
        return parse_xlsx(file_path)
    elif file_type == 'email':
        return parse_email(file_path)
    else:
        return "Unsupported file type"


def chunk_text(text, chunk_size=1000, chunk_overlap=200):
    """
    Split text into chunks with overlap for better context preservation.
    
    Args:
        text: The text to chunk
        chunk_size: Maximum size of each chunk in characters
        chunk_overlap: Number of characters to overlap between chunks
        
    Returns:
        list: List of tuples (chunk_text, start_char, end_char)
    """
    if not text or len(text.strip()) == 0:
        return []
    
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        # Calculate end position
        end = min(start + chunk_size, text_length)
        
        # Try to break at sentence boundaries for better chunk quality
        if end < text_length:
            # Look for sentence endings near the chunk boundary
            for punct in ['. ', '.\n', '! ', '!\n', '? ', '?\n', '\n\n']:
                last_punct = text.rfind(punct, start, end)
                if last_punct != -1 and last_punct > start + chunk_size * 0.7:
                    end = last_punct + len(punct)
                    break
        
        # Extract chunk
        chunk_text = text[start:end].strip()
        if chunk_text:
            chunks.append((chunk_text, start, end))
        
        # Move start position with overlap
        start = max(start + 1, end - chunk_overlap)
    
    return chunks

